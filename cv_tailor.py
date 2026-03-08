"""
cv_tailor.py - Core logic for CV tailoring using Claude API.
Reads a .docx CV, sends it with a JD to Claude, applies the suggested edits,
and saves the result as .docx and PDF.
"""

import json
import platform
import shutil
import subprocess
from pathlib import Path

import anthropic
import docx
import pdfplumber


def extract_pdf_paragraphs(pdf_path: str) -> list[dict]:
    """Extract text lines from a PDF CV as paragraph dicts."""
    paragraphs = []
    idx = 0
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            for line in text.split("\n"):
                line = line.strip()
                if line:
                    paragraphs.append({"index": idx, "text": line, "style": "Normal"})
                    idx += 1
    return paragraphs


def extract_cv_paragraphs(path: str) -> list[dict]:
    """
    Extract all paragraphs from a CV file (.docx or .pdf).
    Returns list of dicts with index, text, and style.
    """
    if path.lower().endswith(".pdf"):
        return extract_pdf_paragraphs(path)
    doc = docx.Document(path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({"index": i, "text": text, "style": para.style.name})
    return paragraphs


def create_docx_from_text(paragraphs: list[dict], docx_output_path: str):
    """Create a plain DOCX from extracted text paragraphs (used when user uploads PDF only)."""
    doc = docx.Document()
    for p in paragraphs:
        doc.add_paragraph(p["text"])
    doc.save(docx_output_path)


def build_prompt(jd_text: str, paragraphs: list[dict]) -> str:
    para_list = "\n".join(
        f"[{p['index']}] ({p['style']}) {p['text']}" for p in paragraphs
    )
    return f"""You are an expert CV writer helping tailor a CV for a specific job.

Below is a job description (JD) followed by the full CV with paragraph indices.

Your task:
1. Identify up to 8 bullet-point paragraphs in the CV that would benefit from being rewritten to better align with this specific JD.
2. For each, provide a rewritten version with:
   - A short bold label (the "category", ending with a colon, e.g. "Customer Success:")
   - A concise plain-text description (the achievement/experience, starting with a space)
3. Only rewrite bullet points (lines that describe experience/achievements). Do NOT touch:
   - Name, contact details, section headers (Education, Experience, etc.)
   - Job titles, company names, dates
   - Skills or education entries that don't need tailoring
4. Preserve factual accuracy - do not fabricate metrics or experiences not present in the original.
5. Keep the tone and length similar to the original bullet points.

Return ONLY a JSON array with this exact structure (no markdown, no explanation):
[
  {{
    "index": <paragraph_index_integer>,
    "original_text": "<exact original text for verification>",
    "bold_label": "<Category Label:>",
    "plain_text": " <achievement description>",
    "reason": "<one sentence explaining why this change helps for this JD>"
  }},
  ...
]

JOB DESCRIPTION:
{jd_text}

CV PARAGRAPHS:
{para_list}
"""


def get_tailoring_suggestions(
    jd_text: str, paragraphs: list[dict], api_key: str
) -> list[dict]:
    """Call Claude API and get suggested CV edits as a list of change dicts."""
    client = anthropic.Anthropic(api_key=api_key)
    prompt = build_prompt(jd_text, paragraphs)

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text.strip()

    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    suggestions = json.loads(raw)
    return suggestions


def set_para(para, bold_text: str, plain_text: str):
    """
    Replace paragraph content preserving font/size from existing runs.
    First run becomes bold label, second run becomes plain description.
    """
    if not para.runs:
        return

    r0 = para.runs[0]
    font_name = r0.font.name
    font_size = r0.font.size
    try:
        font_color = r0.font.color.rgb if r0.font.color and r0.font.color.type else None
    except Exception:
        font_color = None

    # Clear all runs
    for run in para.runs:
        run.text = ""

    # Bold label in run 0
    r0.text = bold_text
    r0.bold = True

    # Plain text in run 1 (create if needed)
    if len(para.runs) >= 2:
        r1 = para.runs[1]
        for run in para.runs[2:]:
            run.text = ""
        r1.text = plain_text
        r1.bold = False
    else:
        new_run = para.add_run(plain_text)
        new_run.bold = False
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size
        if font_color:
            try:
                new_run.font.color.rgb = font_color
            except Exception:
                pass


def apply_suggestions(
    src_path: str, dst_path: str, suggestions: list[dict]
) -> list[dict]:
    """
    Copy CV, apply Claude's suggested edits, save to dst_path.
    Returns list of applied changes with verification info.
    """
    shutil.copy2(src_path, dst_path)
    doc = docx.Document(dst_path)
    paras = doc.paragraphs

    applied = []
    skipped = []

    for s in suggestions:
        idx = s["index"]
        if idx >= len(paras):
            skipped.append({"index": idx, "reason": "Index out of range"})
            continue

        para = paras[idx]
        actual_text = para.text.strip()
        expected_text = s.get("original_text", "").strip()

        # Verify text matches (fuzzy - check if key words overlap)
        if expected_text and actual_text != expected_text:
            # Try to find by text match if index shifted
            match_idx = None
            for j, p in enumerate(paras):
                if p.text.strip() == expected_text:
                    match_idx = j
                    break
            if match_idx is not None:
                idx = match_idx
                para = paras[idx]
            else:
                skipped.append(
                    {
                        "index": idx,
                        "original_text": expected_text,
                        "actual_text": actual_text,
                        "reason": "Text mismatch - paragraph may have shifted",
                    }
                )
                continue

        set_para(para, s["bold_label"], s["plain_text"])
        applied.append(
            {
                "index": idx,
                "original_text": actual_text,
                "new_label": s["bold_label"],
                "new_text": s["plain_text"],
                "reason": s.get("reason", ""),
            }
        )

    doc.save(dst_path)
    return applied, skipped


def generate_pdf(docx_path: str, pdf_path: str):
    """Convert .docx to PDF. Uses docx2pdf on Windows, LibreOffice on Linux."""
    if platform.system() == "Windows":
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    else:
        out_dir = str(Path(pdf_path).parent)
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True, capture_output=True, timeout=60,
        )
        # LibreOffice names the output after the input stem; rename if needed
        generated = Path(out_dir) / (Path(docx_path).stem + ".pdf")
        if generated != Path(pdf_path):
            generated.rename(pdf_path)


def tailor_cv(
    template_path: str,
    jd_text: str,
    output_stem: str,
    output_dir: str,
    api_key: str,
) -> dict:
    """
    Full pipeline: extract paragraphs -> get suggestions -> apply -> save docx + pdf.

    Returns dict with:
      - docx_path: path to output .docx
      - pdf_path: path to output .pdf
      - applied: list of applied changes
      - skipped: list of skipped changes
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = str(output_dir / f"{output_stem}.docx")
    pdf_path = str(output_dir / f"{output_stem}.pdf")

    # Step 1: Extract paragraphs
    paragraphs = extract_cv_paragraphs(template_path)

    # Step 2: Get Claude's suggestions
    suggestions = get_tailoring_suggestions(jd_text, paragraphs, api_key)

    # Step 3: Apply suggestions
    applied, skipped = apply_suggestions(template_path, docx_path, suggestions)

    # Step 4: Convert to PDF
    generate_pdf(docx_path, pdf_path)

    return {
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "applied": applied,
        "skipped": skipped,
        "suggestions": suggestions,
    }
