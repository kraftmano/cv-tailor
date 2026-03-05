"""
build_templates.py
Generates the four CV templates in cv_tailor_app/templates/ from source CVs.
Run once to (re)build all templates.

Templates:
  product.docx        <- _Oliver_Kraftman_CV_vF2_product.docx  (already tailored)
  customer_success.docx <- Decagon/_Oliver_Kraftman_CV_Decagon.docx (already tailored)
  growth.docx         <- _Oliver_Kraftman_CV_vF2.docx + growth rewrites
  chief_of_staff.docx <- _Oliver_Kraftman_CV_vF2.docx + CoS rewrites
"""

import shutil
from pathlib import Path
import docx

CV_DIR = Path(r"c:\Users\User\Dropbox\Personal Files\INSEAD\5_Career\CV")
TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

BASE      = CV_DIR / "_Oliver_Kraftman_CV_vF2.docx"
PRODUCT   = CV_DIR / "_Oliver_Kraftman_CV_vF2_product.docx"
DECAGON   = CV_DIR / "Decagon" / "_Oliver_Kraftman_CV_Decagon.docx"


def set_para(para, bold_text: str, plain_text: str):
    """Replace paragraph content, preserving font/size from existing runs."""
    if not para.runs:
        return
    r0 = para.runs[0]
    font_name = r0.font.name
    font_size = r0.font.size
    try:
        font_color = r0.font.color.rgb if r0.font.color and r0.font.color.type else None
    except Exception:
        font_color = None

    for run in para.runs:
        run.text = ""

    r0.text = bold_text
    r0.bold = True

    if len(para.runs) >= 2:
        r1 = para.runs[1]
        for run in para.runs[2:]:
            run.text = ""
        r1.text = plain_text
        r1.bold = False
    else:
        new_run = para.add_run(plain_text)
        new_run.bold = False
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size
        if font_color:
            try:
                new_run.font.color.rgb = font_color
            except Exception:
                pass


# ── 1. PRODUCT ────────────────────────────────────────────────────────────────
dst = TEMPLATES_DIR / "product.docx"
shutil.copy2(PRODUCT, dst)
print(f"product.docx  <- {PRODUCT.name}")


# ── 2. CUSTOMER SUCCESS ───────────────────────────────────────────────────────
dst = TEMPLATES_DIR / "customer_success.docx"
shutil.copy2(DECAGON, dst)
print(f"customer_success.docx  <- {DECAGON.name}")


# ── 3. GROWTH ─────────────────────────────────────────────────────────────────
# Base: vF2 paragraphs (indices from base CV)
# [22] Operations          -> Growth Partnerships
# [28] Business Development -> Revenue Growth
# [29] Strategy            -> Strategic Partnerships & GTM
# [35] Strategy & Execution -> User Acquisition
# [36] Product Development  -> Growth Engineering
# [42] Marketing           -> Audience Growth & Brand
# [53] Coding Languages    -> Growth Tools & Analytics

dst = TEMPLATES_DIR / "growth.docx"
shutil.copy2(BASE, dst)
doc = docx.Document(dst)
p = doc.paragraphs

# Hakira – reframe around growth/BD
set_para(p[22],
    "Growth Partnerships:",
    " Built proprietary CRM and outreach infrastructure from scratch, driving a 50% uplift in re-engagements and securing 3 new deal mandates through data-driven relationship management")

# Winkreative – BD and revenue growth to the fore
set_para(p[28],
    "Revenue Growth:",
    " Up-sold \u00a3500K of additional services and achieved 80% client renewal and extension rate by consistently delivering measurable value and deepening strategic partnerships")

set_para(p[29],
    "GTM & Market Intelligence:",
    " Authored sector briefings for CEO meetings with prospective clients; supported new market entry strategies that enabled successful pitches and expanded Winkreative\u2019s client base")

# Collate.org – growth and acquisition narrative
set_para(p[35],
    "User Acquisition:",
    " Scaled platform to 2,500+ users on near-zero budget via automated outreach tooling, a UK-wide content competition with 100+ entries, and partnerships with university media societies")

set_para(p[36],
    "Growth Engineering:",
    " Built and shipped a web platform for \u00a350K, achieving 10x SEO improvement and 90% faster page load times, directly compounding organic user acquisition")

# The Broad – audience growth
set_para(p[42],
    "Audience Growth:",
    " Scaled readership to 3,000 unique readers/month through targeted on-campus activation; led team of 5 to produce a 200+ attendee launch event and two panel discussions with 700+ attendees")

# Additional information
set_para(p[53],
    "Growth & Analytics Tools:",
    " Python, STATA; CRM implementation and automated outreach; AI/ML coursework including Hands-on Deep Learning and agentic AI exploration")

doc.save(dst)
print(f"growth.docx   <- {BASE.name} + rewrites")


# ── 4. CHIEF OF STAFF ─────────────────────────────────────────────────────────
# [20] Pitch Documentation    -> Executive Communications
# [21] Deal Execution         -> Cross-Functional Deal Coordination
# [22] Operations             -> Operations & Systems Design
# [27] Leadership & Financial -> Programme Management
# [29] Strategy               -> Strategic Research & Briefing
# [33] Financing (Collate)    -> Founding Operations (reorder to lead at Collate)
# [34] Leadership (Collate)   -> Team Building & Governance
# [35] Strategy & Execution   -> Strategic Execution
# [52] Financial Analysis     -> keep but tweak
# [53] Coding Languages       -> Operations & Analytical Tools

dst = TEMPLATES_DIR / "chief_of_staff.docx"
shutil.copy2(BASE, dst)
doc = docx.Document(dst)
p = doc.paragraphs

# Hakira – operations and exec comms
set_para(p[20],
    "Executive Communications:",
    " Authored fundraising materials and executive briefings for senior stakeholders; led documentation for VC third-fund raise achieving oversubscribed \u00a340M close")

set_para(p[21],
    "Cross-Functional Coordination:",
    " Brokered introductions to 4 partner firms and co-ordinated multi-party deal teams across legal, finance and advisory functions to facilitate a \u00a313M cash-shell IPO")

set_para(p[22],
    "Operations & Systems Design:",
    " Designed and implemented a company-wide CRM and document management system from scratch, driving a 50% increase in relationship re-engagements and 3 new mandates")

# Winkreative – programme management, P&L, strategic advisory
set_para(p[27],
    "Programme Management:",
    " Managed budgets up to \u00a3800K and led cross-functional teams of 6\u20138 across design, tech and strategy to deliver complex multi-channel client programmes on time and within brief")

set_para(p[29],
    "Strategic Research & Briefing:",
    " Authored C-suite briefing documents and conducted sector research for strategy director, enabling high-value pitches and supporting client consultancy engagements")

# Collate.org – founding ops and multi-stakeholder
set_para(p[33],
    "Founding Operations:",
    " Secured \u00a3155K angel investment, designed company operating model, and managed financial controls and board reporting as sole founding operator")

set_para(p[34],
    "Team Building & Governance:",
    " Recruited and managed 5 interns and assembled a board of advisors from tech, media and publishing; established governance and reporting structures for a scaling organisation")

set_para(p[35],
    "Strategic Execution:",
    " Translated vision into operational roadmap; coordinated product, marketing and partnerships functions to scale platform to 2,500+ users with minimal resources")

# Additional information
set_para(p[53],
    "Operations & Analytical Tools:",
    " Python, STATA; CRM systems, financial modelling (Wall Street Prep); AI/ML coursework including Hands-on Deep Learning and agentic AI tool exploration")

doc.save(dst)
print(f"chief_of_staff.docx <- {BASE.name} + rewrites")

print("\nAll templates built successfully.")
print(f"Location: {TEMPLATES_DIR}")
